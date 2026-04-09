"""Rich TUI command-line interface for CX_DB8."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Annotated, Optional

import typer
from rich import box
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Confirm, FloatPrompt, Prompt
from rich.table import Table
from rich.text import Text

from cx_db8 import __version__
from cx_db8.display import export_html, export_svg, print_summary
from cx_db8.embeddings import DEFAULT_MODEL, Embedder
from cx_db8.summarizer import Granularity, summarize

app = typer.Typer(
    name="cx-db8",
    help="Unsupervised contextual extractive summarizer for debate evidence.",
    no_args_is_help=True,
    rich_markup_mode="rich",
)
console = Console()


BANNER = r"""
[bold magenta] ██████╗██╗  ██╗     ██████╗ ██████╗  █████╗ [/]
[bold magenta]██╔════╝╚██╗██╔╝     ██╔══██╗██╔══██╗██╔══██╗[/]
[bold magenta]██║      ╚███╔╝      ██║  ██║██████╔╝ █████╔╝[/]
[bold magenta]██║      ██╔██╗      ██║  ██║██╔══██╗██╔══██╗[/]
[bold magenta]╚██████╗██╔╝ ██╗     ██████╔╝██████╔╝╚█████╔╝[/]
[bold magenta] ╚═════╝╚═╝  ╚═╝     ╚═════╝ ╚═════╝  ╚════╝ [/]
[dim]Unsupervised Contextual Extractive Summarizer v{version}[/]
"""


def _read_card_text(file: Path | None) -> str:
    """Read card text from file, stdin, or interactive prompt."""
    if file is not None:
        return file.read_text(encoding="utf-8")

    if not sys.stdin.isatty():
        return sys.stdin.read()

    console.print(
        "[dim]Paste card text below, then press [bold]Ctrl-D[/bold] "
        "(Ctrl-Z + Enter on Windows) to finish:[/]"
    )
    lines = []
    try:
        while True:
            lines.append(input())
    except EOFError:
        pass
    return "\n".join(lines)


def _show_banner() -> None:
    console.print(BANNER.format(version=__version__))


@app.command()
def run(
    file: Annotated[
        Optional[Path],
        typer.Option("--file", "-f", help="Path to a text file containing the card/evidence."),
    ] = None,
    query: Annotated[
        Optional[str],
        typer.Option("--query", "-q", help="Card tag / query. Use -1 for generic summary."),
    ] = None,
    granularity: Annotated[
        Granularity,
        typer.Option("--granularity", "-g", help="Summarization granularity."),
    ] = Granularity.SENTENCE,
    underline: Annotated[
        Optional[float],
        typer.Option("--underline", "-u", help="Underline percentile (1-99)."),
    ] = None,
    highlight: Annotated[
        Optional[float],
        typer.Option("--highlight", "-H", help="Highlight percentile (1-99)."),
    ] = None,
    model: Annotated[
        str,
        typer.Option("--model", "-m", help="Sentence-transformer model name."),
    ] = DEFAULT_MODEL,
    word_window: Annotated[
        int,
        typer.Option("--word-window", "-w", help="Word/phrase-level context window size."),
    ] = 10,
    bridge_gap: Annotated[
        int,
        typer.Option("--bridge-gap", "-b", help="Max gap to bridge in phrase mode (keeps grammar)."),
    ] = 3,
    output_docx: Annotated[
        Optional[Path],
        typer.Option("--docx", help="Export summary as Word document."),
    ] = None,
    output_html: Annotated[
        Optional[Path],
        typer.Option("--html", help="Export summary as HTML file."),
    ] = None,
    output_svg: Annotated[
        Optional[Path],
        typer.Option("--svg", help="Export summary as SVG screenshot."),
    ] = None,
    visualize: Annotated[
        bool,
        typer.Option("--viz", help="Show 3D embedding visualization (requires viz extra)."),
    ] = False,
    interactive: Annotated[
        bool,
        typer.Option("--interactive", "-i", help="Run in interactive loop mode."),
    ] = False,
) -> None:
    """Summarize debate evidence using contextual embeddings.

    Supports word, sentence, and paragraph-level extractive summarization
    with customizable thresholds and multiple output formats.
    """
    _show_banner()

    if interactive:
        _interactive_loop(model, granularity, word_window)
        return

    # Load model
    with console.status("[bold cyan]Loading embedding model...", spinner="dots"):
        embedder = Embedder(model)

    # Read card text
    card_text = _read_card_text(file)
    if not card_text.strip():
        console.print("[red]Error: No card text provided.[/]")
        raise typer.Exit(1)

    # Get query
    if query is None:
        query = Prompt.ask(
            "[bold cyan]Card tag / query[/] [dim](-1 for generic summary)[/]"
        )
    if query == "-1":
        query = ""

    # Get percentiles
    if underline is None:
        underline = FloatPrompt.ask(
            "[bold cyan]Underline percentile[/] [dim](1-99, e.g. 70 = top 30%)[/]",
            default=70.0,
        )
    if highlight is None:
        highlight = FloatPrompt.ask(
            "[bold cyan]Highlight percentile[/] [dim](should be > underline)[/]",
            default=85.0,
        )

    # Summarize
    with console.status("[bold cyan]Computing embeddings & summarizing...", spinner="dots"):
        result = summarize(
            card_text=card_text,
            query=query,
            embedder=embedder,
            granularity=granularity,
            underline_percentile=underline,
            highlight_percentile=highlight,
            word_window_size=word_window,
            bridge_gap_size=bridge_gap,
            want_embeddings=visualize,
        )

    # Display
    print_summary(result, console=console)

    # Exports
    if output_docx:
        from cx_db8.docx_export import export_docx

        export_docx(result, output_docx)
        console.print(f"[green]Word document saved to {output_docx}[/]")

    if output_html:
        export_html(result, output_html)
        console.print(f"[green]HTML saved to {output_html}[/]")

    if output_svg:
        export_svg(result, output_svg)
        console.print(f"[green]SVG saved to {output_svg}[/]")

    if visualize:
        from cx_db8.graph import visualize_3d

        visualize_3d(result)


def _interactive_loop(model_name: str, granularity: Granularity, word_window: int) -> None:
    """Run CX_DB8 in a continuous interactive loop."""
    with console.status("[bold cyan]Loading embedding model...", spinner="dots"):
        embedder = Embedder(model_name)

    console.print(
        Panel(
            "[bold]Interactive mode[/]\n"
            "Summarize multiple cards in sequence.\n"
            "Type [bold cyan]quit[/] or press [bold]Ctrl-C[/] to exit.",
            border_style="magenta",
        )
    )

    docx_results: list[tuple] = []

    while True:
        console.print()
        if not Confirm.ask("[bold cyan]Summarize a card?[/]", default=True):
            break

        # Card text
        console.print(
            "[dim]Paste card text, then press [bold]Ctrl-D[/bold] to finish:[/]"
        )
        lines = []
        try:
            while True:
                lines.append(input())
        except EOFError:
            pass
        card_text = "\n".join(lines)

        if not card_text.strip():
            console.print("[yellow]Empty card text, skipping.[/]")
            continue

        query = Prompt.ask(
            "[bold cyan]Card tag / query[/] [dim](-1 for generic)[/]"
        )
        if query == "-1":
            query = ""

        underline = FloatPrompt.ask(
            "[bold cyan]Underline percentile[/]", default=70.0
        )
        highlight = FloatPrompt.ask(
            "[bold cyan]Highlight percentile[/]", default=85.0
        )
        author = Prompt.ask("[bold cyan]Author & date[/]", default="")
        citation = Prompt.ask("[bold cyan]Citation[/]", default="")

        with console.status("[bold cyan]Summarizing...", spinner="dots"):
            result = summarize(
                card_text=card_text,
                query=query,
                embedder=embedder,
                granularity=granularity,
                underline_percentile=underline,
                highlight_percentile=highlight,
                word_window_size=word_window,
            )

        print_summary(result, author=author, citation=citation, console=console)
        docx_results.append((result, author, citation))

    # Save collected summaries
    if docx_results and Confirm.ask(
        "[bold cyan]Save all summaries to Word document?[/]", default=True
    ):
        from cx_db8.docx_export import export_docx

        path = Path(
            Prompt.ask("[bold cyan]Output path[/]", default="cx_db8_summaries.docx")
        )
        from docx import Document as DocxDocument

        doc = DocxDocument()
        for res, auth, cite in docx_results:
            heading = res.query[:100] if res.query else "Summary"
            doc.add_heading(heading, level=1)
            if auth:
                p = doc.add_paragraph()
                p.add_run(auth).bold = True
            if cite:
                doc.add_paragraph(cite)
            para = doc.add_paragraph()
            for span in res.spans:
                run = para.add_run(span.text + " ")
                if span.score >= res.highlight_threshold:
                    run.bold = True
                    run.underline = True
                elif span.score >= res.underline_threshold:
                    run.underline = True
            doc.add_page_break()
        doc.save(str(path))
        console.print(f"[green]Saved {len(docx_results)} summaries to {path}[/]")


@app.command()
def models() -> None:
    """List recommended sentence-transformer models."""
    _show_banner()
    table = Table(
        title="Recommended Models",
        box=box.ROUNDED,
        title_style="bold cyan",
    )
    table.add_column("Model", style="bold")
    table.add_column("Size", justify="right")
    table.add_column("Speed", justify="center")
    table.add_column("Quality", justify="center")
    table.add_column("Notes")

    table.add_row(
        "all-MiniLM-L6-v2",
        "80 MB",
        "[green]Fast[/]",
        "[yellow]Good[/]",
        "Default. Great balance of speed and quality.",
    )
    table.add_row(
        "all-mpnet-base-v2",
        "420 MB",
        "[yellow]Medium[/]",
        "[green]Best[/]",
        "Highest quality general-purpose model.",
    )
    table.add_row(
        "multi-qa-MiniLM-L6-cos-v1",
        "80 MB",
        "[green]Fast[/]",
        "[yellow]Good[/]",
        "Optimized for semantic search / QA.",
    )
    table.add_row(
        "all-distilroberta-v1",
        "290 MB",
        "[yellow]Medium[/]",
        "[yellow]Good[/]",
        "Good alternative general-purpose model.",
    )
    table.add_row(
        "BAAI/bge-small-en-v1.5",
        "130 MB",
        "[green]Fast[/]",
        "[green]Great[/]",
        "State-of-the-art small model.",
    )

    console.print(table)
    console.print(
        "\n[dim]Use any model from https://huggingface.co/models?library=sentence-transformers[/]"
    )
    console.print("[dim]Pass with: cx-db8 run --model MODEL_NAME[/]\n")


@app.command()
def version() -> None:
    """Show version information."""
    console.print(f"[bold magenta]CX_DB8[/] v{__version__}")


@app.command()
def demo() -> None:
    """Run a built-in demo with sample debate evidence."""
    _show_banner()

    sample_text = (
        "If the evidence gets out to the public while the scientists are still analyzing "
        "the signal, Forgan said they could manage the public's expectations by using "
        "something called the Rio Scale. It's essentially a numeric value that represents "
        "the degree of likelihood that an alien contact is \"real.\" Forgan added that the "
        "Rio Scale is also undergoing an update, and more should be coming out about it in May. "
        "If the aliens did arrive here, \"first contact\" protocols likely would be useless, "
        "because if they're smart enough to show up physically, they could probably do anything "
        "else they like, according to Shostak. \"Personally, I would leave town,\" Shostak "
        "quipped. \"I have no idea what they are here for.\" But there's little need to worry. "
        "An \"Independence Day\" scenario of aliens blowing up important national buildings such "
        "as the White House is extremely unlikely, Forgan said, because interstellar travel is "
        "difficult. Early SETI work: To find a signal, first we have to be listening for it. "
        "SETI \"listening\" is going on all over the world, and in fact, this has been happening "
        "for many decades. The first modern SETI experiment took place in 1960. Under Project "
        "Ozma, Cornell University astronomer Frank Drake pointed a radio telescope located at "
        "Green Bank, West Virginia at two stars called Tau Ceti and Epsilon Eridani. He scanned "
        "at a frequency astronomers nickname \"the water hole,\" which is close to the frequency "
        "of light that's given off by hydrogen and hydroxyl. In 1977, The Ohio State University "
        "SETI's program made international headlines after a project volunteer, Jerry Ehman, "
        "wrote, \"Wow!\" beside a strong signal a telescope there received. The Aug. 15, 1977, "
        "\"Wow\" signal was never repeated, however."
    )
    sample_query = "SETI signal detection and the Wow Signal"

    console.print(
        Panel(
            "[bold]Demo Mode[/]\n"
            "Running summarization on sample SETI evidence...",
            border_style="cyan",
        )
    )
    console.print()

    with console.status("[bold cyan]Loading embedding model...", spinner="dots"):
        embedder = Embedder()

    for gran in [Granularity.SENTENCE, Granularity.PHRASE]:
        with console.status(
            f"[bold cyan]Summarizing at {gran.value} level...", spinner="dots"
        ):
            result = summarize(
                card_text=sample_text,
                query=sample_query,
                embedder=embedder,
                granularity=gran,
                underline_percentile=60.0,
                highlight_percentile=80.0,
            )
        print_summary(result, author="Ehman 77", console=console)

    console.print("[bold green]Demo complete![/]\n")


if __name__ == "__main__":
    app()
