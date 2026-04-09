"""Word document export for CX_DB8 summaries."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from cx_db8.summarizer import SummaryResult


def export_docx(
    result: SummaryResult,
    path: Path,
    author: str = "",
    citation: str = "",
) -> Path:
    """Export a summary as a formatted Word document."""
    doc = Document()

    # Title
    heading = result.query[:100] if result.query else "Summary"
    doc.add_heading(heading, level=1)

    # Author / citation
    if author:
        p = doc.add_paragraph()
        run = p.add_run(author)
        run.bold = True
    if citation:
        doc.add_paragraph(citation)

    # Summary body
    para = doc.add_paragraph()
    for span in result.spans:
        run = para.add_run(span.text + " ")
        if span.score >= result.highlight_threshold:
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0x8B, 0x6C, 0x00)  # dark gold
        elif span.score >= result.underline_threshold:
            run.underline = True
        # else: plain text, no formatting

    doc.save(str(path))
    return path
