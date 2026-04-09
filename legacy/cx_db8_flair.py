import argparse
import flair
from flair.data import Sentence
from flair.models import SequenceTagger
from flair.embeddings import WordEmbeddings, FlairEmbeddings, StackedEmbeddings, DocumentPoolEmbeddings, BertEmbeddings, ELMoEmbeddings, OpenAIGPTEmbeddings, TransformerXLEmbeddings, XLNetEmbeddings
import torch
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.metrics import jaccard_similarity_score
#import numpy as np
from docx import Document
import sys
import numpy as np
from itertools import islice
from collections import deque
import matplotlib 
#import umap ##only necessary for weird experiments
import matplotlib.pyplot as plt
#import seaborn as sns ##only necessary for weird experiments
from mpl_toolkits.mplot3d import proj3d
import matplotlib.cm as cm
from torch.nn import CosineSimilarity
from sty import fg, bg, ef, rs, RgbFg
from sklearn.preprocessing import MinMaxScaler
import syntok.segmenter as segmenter
from ansi2html import Ansi2HTMLConverter
import text_graph 

conv = Ansi2HTMLConverter()

document = Document() ## Create a python-docx document
cos = CosineSimilarity(dim=1, eps=1e-6)



class bcolors: #For converting the ANSI string to HTML - Sty is not supported well :(
    HIGHLIGHT = '\33[43m'
    END = '\033[0m'


granularity_level = "Sent" #"Word" "Sent" "Paragraph"
dynamic = False  ##Controls if we highlight the more important words more or not
graph = False ###ONLY WORKS WITH "Word" granularity_level
word_doc = True 
html = True
word_window_size = 10 ##This is really doubled since it's Bi-directional. Only matters for word level granularity
highlight_color_intensity = 175 # try values between 25 and 200
doc_embeddings = []
scores = []

stacked_embeddings = DocumentPoolEmbeddings([
                                        #WordEmbeddings('en'),
                                        XLNetEmbeddings('xlnet-large-cased', use_scalar_mix = 'True')
                                        #WordEmbeddings('glove'),
                                        #WordEmbeddings('extvec'),#ELMoEmbeddings('original'),
                                        #BertEmbeddings('bert-base-cased'),
                                        #FlairEmbeddings('news-forward-fast'),
                                        #FlairEmbeddings('news-backward-fast'),
                                        #OpenAIGPTEmbeddings()
                                        #TransformerXLEmbeddings()
                                        ]) #, mode='max')

def arg_parser():
  parser = argparse.ArgumentParser()
  parser.add_argument("--card_file_path", "-c", type=str,
                      help="Path to file for file text")
  args = parser.parse_args()
  return args

def set_card_text(card_path=None):
  card = ""
  if card_path:
    with open(card_path, "r") as in_file:
      card = in_file.read()
  else:
    print("Input the Card Text, press Ctrl-D to end text entry")
    card = sys.stdin.read()#input("Input the Card Text: ")
  return card


def set_card(card_path=None):
    card = set_card_text(card_path)
    card_tag = input("Input the card_tag, or a -1 to summarize in-terms of the card itself: ")
    card = str(card)
    if str(card_tag) == "-1": #This will not work with large documents when bert is enabled
        card_tag = Sentence(str(card))
        tag_str = ""
    else:
        tag_str = str(card_tag)
        card_tag = Sentence(str(card_tag))
    return card, card_tag, tag_str


def create_ngram(num_context, card, card_tag):
    #card = set_card()
    card_words_org = card.split()
    ngram_list = []
    #print(len(card_words_org))

    for i in range(0, len(card_words_org)):
        # Get the sliding window.
        lower_bound = i - num_context if i - num_context > 0 else 0
        upper_bound = i + num_context if i + num_context < len(card_words_org) else len(card_words_org) - 1
        new_word = card_words_org[lower_bound : upper_bound]
        print(new_word)

        # Join the window.
        new_string = " ".join(new_word)
        if new_string == "":
          new_string = " "
        ngram_list.append(Sentence(new_string))

    return ngram_list, card_words_org


def embed(card_tag, card_as_sentence, card_words, card_words_org):
    stacked_embeddings.embed(card_tag)
    #stacked_embeddings.embed(card_as_sentence)
    #print(card_as_sentence.get_embedding().reshape(1,-1))
    word_list = []
    token_removed_ct = 0
    card_tag_emb = card_tag.get_embedding()
    if granularity_level == "Word":
        for word, count in zip(card_words_org, range(0, len(card_words_org))):
            n_gram_word = card_words[count]
            stacked_embeddings.embed(n_gram_word)
            n_gram_emb = n_gram_word.get_embedding()
            if graph:
                doc_embeddings.append(n_gram_emb.cpu().detach().numpy())
            word_sim = cos(card_tag_emb.reshape(1,-1), n_gram_emb.reshape(1, -1))
            word_tup = (card_words_org[count], word_sim) #card_words_org[count]
            word_list.append(word_tup)
        if graph:
            doc_embeddings.append(card_tag_emb.cpu().detach().numpy())
        print(len(word_list))
        print(len(card_words))
        print(len(card_words_org))
    else: 
        for sentence in card_as_sentence:
            set_obj = Sentence(sentence)
            stacked_embeddings.embed(set_obj)
            sentence_emb = set_obj.get_embedding()
            word_sim = cos(card_tag_emb.reshape(1,-1), sentence_emb.reshape(1, -1))
            sentence_tup = (sentence, word_sim)
            word_list.append(sentence_tup)
    return word_list


def summarize(word_list):
    return [sum_word[1] for sum_word in word_list]

def get_sim_scores(word_list):
    return [float(sum_word[1]) for sum_word in word_list]

def run_loop(context, card, card_tag):
    list_of_sentences = []
    list_of_paragraphs = []
    if granularity_level == "Sent":
        for paragraph in segmenter.analyze(card):
            for sentence in paragraph: ## sentence level summarization
                set_str = ""
                for token in sentence:
                    set_str += token.spacing
                    set_str += token.value
                list_of_sentences.append(set_str)
        word_list = embed(card_tag, list_of_sentences, 0, 0)
    elif granularity_level == "Paragraph":
        for paragraph in segmenter.analyze(card):
            set_str = ""
            for sentence in paragraph: ## sentence level summarization
                #set_str = ""
                for token in sentence:
                    set_str += token.spacing
                    set_str += token.value
            list_of_paragraphs.append(set_str)
        word_list = embed(card_tag, list_of_paragraphs, 0, 0)
    elif granularity_level == "Word":
        card_as_sentence = Sentence(card)
        card_words, card_words_org = create_ngram(context, card, card_tag)
        word_list = embed(card_tag, card_as_sentence, card_words, card_words_org)
    #print(word_list)
    return word_list



def parse_word_val_list(word_list, h, n, val_list):
    sum_str = ""
    html_str = ""
    removed_str = ""
    token_removed_ct = 0
    to_highlight = np.asarray([i for i in val_list if i > h])
    scaler = MinMaxScaler(feature_range = (20, 255))
    scaler.fit(to_highlight.reshape(-1, 1))
    for sum_word in word_list:
        prev_word = "n"
        if float(sum_word[1]) > h:
            if dynamic:
                bgnum = int(scaler.transform(sum_word[1].cpu().detach().numpy().reshape(-1, 1)))
            else:
                bgnum = highlight_color_intensity
            html_str += ef.u + bcolors.HIGHLIGHT + sum_word[0] + " "
            sum_str += ef.u + bg(bgnum, bgnum, 0)
            runner = par.add_run(sum_word[0] + " ")
            runner.underline = True
            runner.bold = True
            sum_str += " ".join([str(sum_word[0]), rs.all])
        elif float(sum_word[1]) > n:
            html_str += ef.u + sum_word[0] + " "
            sum_str += ef.u
            runner = par.add_run(sum_word[0] + " ")
            runner.underline = True
            sum_str += " ".join([str(sum_word[0]), rs.all])
        else:
            html_str += sum_word[0] + " " + bcolors.END
            token_removed_ct += 1
            sum_str += str(sum_word[0])
            runner = par.add_run(sum_word[0] + " ")
            sum_str += " "
            removed_str += str(sum_word[0])
            removed_str += " "
    return sum_str, removed_str, token_removed_ct, html_str

def new_sum(word_list, underline_p, highlight_p):
    val_list = summarize(word_list)
    th_underline = np.percentile(val_list, float(underline_p))
    th_highlight = np.percentile(val_list, float(highlight_p))
    sum_str, removed_str, token_removed_ct, t_html_str = parse_word_val_list(word_list, th_highlight, th_underline, val_list)


    print("CARD: ")
    print(card)
    print("CARD_TAG :")
    print(card_tag)
    print("GENERATED SUMMARY: ")
    print(sum_str)
    print("tokens removed:" + " " + str(token_removed_ct))
    #print(conv.convert(sum_str))
    if html:
        #print(conv.convert(t_html_str))
        with open("cx_html.html", "a") as hf:
            hf.write(conv.convert(t_html_str))





for i in range(0, 1000):
    args = arg_parser()
    cont = input("Do you want to continue? type y for yes else for stop!")
    if cont == "y":
        card, card_tag, tag_str = set_card(args.card_file_path)
        head = document.add_heading(tag_str, level=1)
        underline_p = input("what percentile should be underlined? (numbers from 1 to 99 acceptable, a 90 means only the top 10% is underlined")
        highlight_p = input("what percentile to highlight? (should be higher than the previous value)") 
        card_auth = input("what's the card author and date?")
        card_cite = input("what's the citation?")
        a_par = document.add_paragraph()
        a_par.add_run(card_auth).bold = True
        c_par = document.add_paragraph(card_cite)
        par = document.add_paragraph()
        word_list = run_loop(word_window_size, card, card_tag)
        if graph:
            #plt.figure(figsize=(15,1))
            #sns.heatmap(doc_embeddings[0:20], robust = True, linewidths=0.7, cmap="PiYG")
            #plt.show()
            gph = np.array([doc_embeddings[k][0:3].tolist() for k in range(0, len(doc_embeddings))])
            word_list.append(('CARD_TAG', 1.0))
            scores = get_sim_scores(word_list)
            #scores.append(1.0) #to account for the query string)
            text_graph.visualize3DData(gph, scores, word_list)

            '''
            #umap_emb = reducer.fit_transform(doc_embeddings) ##if we want to fit with the card tag, maybe not as interesting
            gph = np.array([doc_embeddings[k][0:3].tolist() for k in range(0, len(doc_embeddings))])
            #umap_emb = reducer.fit_transform((doc_embeddings[:-1]))
            #t_emb = reducer.transform(doc_embeddings[-1].reshape(1, -1))
            #n_umap_emb = np.append(umap_emb, t_emb, axis=0)
            scores = get_sim_scores(word_list)
            scores.append(1.0) #to account for the query string)
            '''
        new_sum(word_list, underline_p, highlight_p)
    else:
        if word_doc:
            document.save('test_sum.docx')
        break


