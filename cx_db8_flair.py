import flair
from flair.data import Sentence
from flair.models import SequenceTagger
from flair.embeddings import WordEmbeddings, FlairEmbeddings, StackedEmbeddings, DocumentPoolEmbeddings, BertEmbeddings, ELMoEmbeddings, OpenAIGPTEmbeddings, TransformerXLEmbeddings
import torch
# create a StackedEmbedding object that combines glove and forward/backward flair embeddings
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.metrics import jaccard_similarity_score
#import numpy as np
from docx import Document
import sys
import numpy as np
from itertools import islice
from collections import deque
import matplotlib 
import umap
import matplotlib.pyplot as plt
import seaborn as sns
from mpl_toolkits.mplot3d import proj3d
import matplotlib.cm as cm
from torch.nn import CosineSimilarity
from sty import fg, bg, ef, rs, RgbFg
from sklearn.preprocessing import MinMaxScaler

document = Document() ## Create a python-docx document
cos = CosineSimilarity(dim=1, eps=1e-6)




dynamic = True
graph = False
doc_embeddings = []
scores = []

stacked_embeddings = DocumentPoolEmbeddings([
                                        #WordEmbeddings('en'),
                                        #WordEmbeddings('glove'),
                                        #WordEmbeddings('extvec'),#ELMoEmbeddings('original'),
                                        #BertEmbeddings('bert-base-cased'),
                                        #FlairEmbeddings('news-forward-fast'),
                                        #FlairEmbeddings('news-backward-fast'),
                                        #OpenAIGPTEmbeddings()
                                        TransformerXLEmbeddings()
                                        ]) #, mode='max')


def set_card():
    print("Input the Card Text, press Ctrl-D to end text entry")
    card = sys.stdin.read()#input("Input the Card Text: ")
    card_tag = input("Input the card_tag, or a -1 to summarize in-terms of the card itself: ")
    card = str(card)
    if str(card_tag) == "-1": #This will not work with large documents when bert is enabled
        card_tag = Sentence(str(card))
        tag_str = ""
    else:
        tag_str = str(card_tag)
        card_tag = Sentence(str(card_tag))
    return card, card_tag, tag_str

def window(seq, n):
    "Returns a sliding window (of width n) over data from the iterable"
    "   s -> (s0,s1,...s[n-1]), (s1,s2,...,sn), ...                   "
    it = iter(seq)
    result = tuple(islice(it, n))
    if len(result) == n:
        yield result
    for elem in it:
        result = result[1:] + (elem,)
        yield result

def create_ngram(num_context, card, card_tag):
    #card = set_card()
    card_words_org = card.split()
    ngram_list = []
    #print(len(card_words_org))

    for i in range(0, len(card_words_org)):
        # Get the sliding window.
        lower_bound = i - num_context if i - num_context > 0 else 0
        upper_bound = i + num_context if i + num_context < len(card_words_org) else len(card_words_org) - 1
        new_word = card_words_org[lower_bound : upper_bound]
        print(new_word)

        # Join the window.
        new_string = " ".join(new_word)
        if new_string == "":
          new_string = " "
        ngram_list.append(Sentence(new_string))

    return ngram_list, card_words_org


def embed(card_tag, card_as_sentence, card_words, card_words_org):
    stacked_embeddings.embed(card_tag)
    #stacked_embeddings.embed(card_as_sentence)
    #print(card_as_sentence.get_embedding().reshape(1,-1))
    word_list = []
    token_removed_ct = 0
    card_tag_emb = card_tag.get_embedding()
    for word, count in zip(card_words_org, range(0, len(card_words_org))):
        n_gram_word = card_words[count]
        stacked_embeddings.embed(n_gram_word)
        n_gram_emb = n_gram_word.get_embedding()
        if graph:
            doc_embeddings.append(n_gram_emb.numpy())
        word_sim = cos(card_tag_emb.reshape(1,-1), n_gram_emb.reshape(1, -1))
        word_tup = (card_words_org[count], word_sim) #card_words_org[count]
        word_list.append(word_tup)
    if graph:
        doc_embeddings.append(card_tag_emb.numpy())
    print(len(word_list))
    print(len(card_words))
    print(len(card_words_org))
    return word_list


def summarize(word_list):
    return [sum_word[1] for sum_word in word_list]

def get_sim_scores(word_list):
    return [float(sum_word[1]) for sum_word in word_list]

def run_loop(context, card, card_tag):
    card_as_sentence = Sentence(card)
    card_words, card_words_org = create_ngram(context, card, card_tag)
    word_list = embed(card_tag, card_as_sentence, card_words, card_words_org)
    return word_list

def parse_word_val_list(word_list, h, n, val_list):
    sum_str = ""
    removed_str = ""
    token_removed_ct = 0
    to_highlight = np.asarray([i for i in val_list if i > h])
    scaler = MinMaxScaler(feature_range = (20, 255))
    scaler.fit(to_highlight.reshape(-1, 1))
    for sum_word in word_list:
        if float(sum_word[1]) > h:
            if dynamic:
                bgnum = int(scaler.transform(sum_word[1].detach().numpy().reshape(-1, 1)))
            else:
                bgnum = 50
            sum_str += ef.u + bg(bgnum, bgnum, 0)
            runner = par.add_run(sum_word[0] + " ")
            runner.underline = True
            runner.bold = True
            sum_str += " ".join([str(sum_word[0]), rs.all])
        elif float(sum_word[1]) > n:
            sum_str += ef.u
            runner = par.add_run(sum_word[0] + " ")
            runner.underline = True
            sum_str += " ".join([str(sum_word[0]), rs.all])
        else:
            token_removed_ct += 1
            sum_str += str(sum_word[0])
            runner = par.add_run(sum_word[0] + " ")
            sum_str += " "
            removed_str += str(sum_word[0])
            removed_str += " "
    return sum_str, removed_str, token_removed_ct

def new_sum(word_list, underline_p, highlight_p):
    val_list = summarize(word_list)
    th_underline = np.percentile(val_list, float(underline_p))
    th_highlight = np.percentile(val_list, float(highlight_p))
    sum_str, removed_str, token_removed_ct = parse_word_val_list(word_list, th_highlight, th_underline, val_list)


    print("CARD: ")
    print(card)
    print("CARD_TAG :")
    print(card_tag)
    print("GENERATED SUMMARY: ")
    print(sum_str)
    print("tokens removed:" + " " + str(token_removed_ct))
    print("NOT UNDERLINED")

    print(removed_str)



"--------------------------------------- Vizualization code below this line ----------------------"

def visualize3DData (X):
    """Visualize data in 3d plot with popover next to mouse position.

    Args:
        X (np.array) - array of points, of shape (numPoints, 3)
    Returns:
        None
    """
    fig = plt.figure(figsize = (16,10))
    ax = fig.add_subplot(111, projection = '3d')
    im = ax.scatter(X[:, 0], X[:, 1], X[:, 2], c = np.asarray(scores), depthshade = False, picker = True, s=30, cmap = "rainbow")
    fig.colorbar(im)


    def distance(point, event):
        """Return distance between mouse position and given data point

        Args:
            point (np.array): np.array of shape (3,), with x,y,z in data coords
            event (MouseEvent): mouse event (which contains mouse position in .x and .xdata)
        Returns:
            distance (np.float64): distance (in screen coords) between mouse pos and data point
        """
        assert point.shape == (3,), "distance: point.shape is wrong: %s, must be (3,)" % point.shape

        # Project 3d data space to 2d data space
        x2, y2, _ = proj3d.proj_transform(point[0], point[1], point[2], plt.gca().get_proj())
        # Convert 2d data space to 2d screen space
        x3, y3 = ax.transData.transform((x2, y2))

        return np.sqrt ((x3 - event.x)**2 + (y3 - event.y)**2)


    def calcClosestDatapoint(X, event):
        """"Calculate which data point is closest to the mouse position.

        Args:
            X (np.array) - array of points, of shape (numPoints, 3)
            event (MouseEvent) - mouse event (containing mouse position)
        Returns:
            smallestIndex (int) - the index (into the array of points X) of the element closest to the mouse position
        """
        distances = [distance (X[i, 0:3], event) for i in range(X.shape[0])]
        return np.argmin(distances)


    def annotatePlot(X, index):
        """Create popover label in 3d chart

        Args:
            X (np.array) - array of points, of shape (numPoints, 3)
            index (int) - index (into points array X) of item which should be printed
        Returns:
            None
        """
        # If we have previously displayed another label, remove it first
        if hasattr(annotatePlot, 'label'):
            annotatePlot.label.remove()
        # Get data point from array of points X, at position index
        x2, y2, _ = proj3d.proj_transform(X[index, 0], X[index, 1], X[index, 2], ax.get_proj())
        annotatePlot.label = plt.annotate(word_list[index][0][0:100],
            xy = (x2, y2), xytext = (-10, 10), textcoords = 'offset points', ha = 'left', va = 'bottom', size = 6,
            bbox = dict(boxstyle = 'round,pad=0.5', fc = 'yellow', alpha = 0.5),
            arrowprops = dict(arrowstyle = '->'))
        fig.canvas.draw()


    def onMouseMotion(event):
        """Event that is triggered when mouse is moved. Shows text annotation over data point closest to mouse."""
        closestIndex = calcClosestDatapoint(X, event)
        annotatePlot (X, closestIndex)

    fig.canvas.mpl_connect('motion_notify_event', onMouseMotion)  # on mouse motion
    plt.show()



for i in range(0, 1000):
    cont = input("Do you want to continue? type y for yes else for stop!")
    if cont == "y":
        card, card_tag, tag_str = set_card()
        head = document.add_heading(tag_str, level=1)
        underline_p = input("what percentile should be underlined? (numbers from 1 to 99 acceptable, a 90 means only the top 10% is underlined")
        highlight_p = input("what percentile to highlight? (should be higher than the previous value)") 
        card_auth = input("what's the card author and date?")
        card_cite = input("what's the citation?")
        a_par = document.add_paragraph()
        a_par.add_run(card_auth).bold = True
        c_par = document.add_paragraph(card_cite)
        par = document.add_paragraph()
        word_list = run_loop(10, card, card_tag)
        if graph:
            plt.figure(figsize=(15,1))
            sns.heatmap(doc_embeddings[0:20], robust = True, linewidths=0.7, cmap="PiYG")
            plt.show()

            '''
            #umap_emb = reducer.fit_transform(doc_embeddings) ##if we want to fit with the card tag, maybe not as interesting
            gph = np.array([doc_embeddings[k][0:3].tolist() for k in range(0, len(doc_embeddings))])
            #umap_emb = reducer.fit_transform((doc_embeddings[:-1]))
            #t_emb = reducer.transform(doc_embeddings[-1].reshape(1, -1))
            #n_umap_emb = np.append(umap_emb, t_emb, axis=0)
            scores = get_sim_scores(word_list)
            scores.append(1.0) #to account for the query string)
            '''
        new_sum(word_list, underline_p, highlight_p)

        if graph:
            '''
            word_list.append(('CARD_TAG', 1.0))
            visualize3DData(gph)
            '''
    else:
        document.save('test_sum.docx')
        break


